"""Provides a class for generating letters for scholarship awards.

Provides the class LetterWriter for generating a letter as
a docx file for a scholarship award. Also provides
the class LetterVariables make it easier to set the values
that will replace the placeholders (variables) in the template letter.
"""

from docx import Document
from io import BytesIO


class LetterVariables:
    """Class for representing Letter Variables.

    Class for representing the variables for
    the placeholders in the template letter.
    """

    def __init__(
        self,
        student_name: str,
        date: str,
        amount: str,
        scholarship_name: str,
        academic_year_fall: str,
        academic_year_spring: str,
        sender_name: str,
        sender_email: str,
        sender_title: str,
    ) -> None:
        """Creates an instance of LetterVariables.

        Creates a new instance of LetterVariables.

        Args:
            student_name (str): Name of the student receiving the award.
            date (str): Date the letter was created.
            amount (str): Amount being awarded.
            scholarship_name (str): Name of the scholarship award.
            academic_year (str): Academic year for the award (E.g, 2023-2024).
            sender_name (str): Name of the sender.
            sender_email (str): Email of the sender.
            sender_title (str): Title / Occupation of the sender.
        """
        self.student_name: str = student_name
        self.date: str = date
        self.amount: str = f"${float(amount):.2f}"
        self.scholarship_name: str = scholarship_name
        self.academic_year: str = f"{academic_year_fall}-{academic_year_spring}"
        self.sender_name: str = sender_name
        self.sender_email: str = sender_email
        self.sender_title: str = sender_title
        self.half_amount: str = f"${(float(amount) / 2):.2f}"

        self.academic_year_fall: str = academic_year_fall
        self.academic_year_spring: str = academic_year_spring

    def __iter__(self):
        """Allows for iterating over attributes and casting to other data structures.

        Allows for iterating over the attributes of the class, as well as
        casting to other data structures.
        """
        yield "{STUDENT_NAME}", self.student_name
        yield "{DATE}", self.date
        yield "{AMOUNT}", self.amount
        yield "{SCHOLARSHIP_NAME}", self.scholarship_name
        yield "{ACADEMIC_YEAR}", self.academic_year
        yield "{SENDER_NAME}", self.sender_name
        yield "{SENDER_EMAIL}", self.sender_email
        yield "{SENDER_TITLE}", self.sender_title
        yield "{HALF_AMOUNT}", self.half_amount
        yield "{ACADEMIC_YEAR_FALL}", self.academic_year_fall
        yield "{ACADEMIC_YEAR_SPRING}", self.academic_year_spring

    def to_dict(self) -> dict:
        """Returns a dict representation.

        Returns a dict representation of the LetterVariables object.

        Returns:
            A dict representing the LetterVariables object.
        """
        return dict(self)

    def __repr__(self) -> str:
        """Returns a str representation.

        Returns a str representation of the LetterVariables object.

        Returns:
            A str representing the LetterVariables object.
        """
        return str(dict(self))


def write_letter(
    template_file_path: str, output_file_path: str, variables: LetterVariables
) -> None:
    """Generates the letter.

    Generates the letter using the template letter and the
    dict containing the keys and values used to replace the
    placeholders in the template letter.
    """
    # Read in the docx file
    document = Document(template_file_path)

    # Convert letter variables to dictionary
    variables_dict: dict[str] = variables.to_dict()

    # Iterate over the items in the dictionary
    for key, value in variables_dict.items():
        for paragraph in document.paragraphs:
            if key in paragraph.text:
                runs = paragraph.runs
                for item in runs:
                    # If the key is present in the text
                    # replace the placeholder with the value
                    if key in item.text:
                        item.text = item.text.replace(key, value)

    # Save the generate docx file to the output file path
    document.save(output_file_path)


def write_letter_to_bytes(template_file_path: str, variables: LetterVariables) -> bytes:
    write_out: BytesIO = BytesIO()
    # Read in the docx file
    document = Document(template_file_path)

    # Convert letter variables to dictionary
    variables_dict: dict[str] = variables.to_dict()

    # Iterate over the items in the dictionary
    for key, value in variables_dict.items():
        for paragraph in document.paragraphs:
            if key in paragraph.text:
                runs = paragraph.runs
                for item in runs:
                    # If the key is present in the text
                    # replace the placeholder with the value
                    if key in item.text:
                        item.text = item.text.replace(key, value)

    # Save the generate docx file to the BytesIO
    document.save(write_out)

    # Get the data from BytesIO as bytes
    document_bytes: bytes = write_out.getvalue()

    return document_bytes


if __name__ == "__main__":
    vars = LetterVariables(
        "Angel Badillo Hernandez",
        "March 8, 2024",
        "1000",
        "Hardworking Scholarship",
        "2024",
        "2025",
        "Angel Badillo",
        "email@email.com",
        "Computer Science Master",
    )

    write_letter("assets/templates/template_letter.docx", "test.docx", vars)

    docx_bytes: bytes = write_letter_to_bytes(
        "assets/templates/template_letter.docx", vars
    )

    with open("test2.docx", "wb") as file:
        file.write(docx_bytes)
